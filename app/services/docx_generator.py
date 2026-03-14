"""Generate ATS-friendly DOCX files from structured CV data.

Uses python-docx to produce clean Word documents that ATS systems can parse
reliably. No tables for layout, no text boxes — just standard headings,
paragraphs, and bullet lists.
"""

import logging
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Emu, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Letter: 8.5 x 11 in; A4: 8.27 x 11.69 in
_LETTER_REGIONS = {"US", "CA", "CO", "VE"}

_PAGE_WIDTH_LETTER = Inches(8.5)
_PAGE_HEIGHT_LETTER = Inches(11)
_PAGE_WIDTH_A4 = Emu(int(8.27 * 914400))
_PAGE_HEIGHT_A4 = Emu(int(11.69 * 914400))
_MARGIN = Inches(1)

# Fonts
_FONT_BODY = "Calibri"
_FONT_HEADING = "Calibri"

# Sizes (Pt)
_SIZE_NAME = 20
_SIZE_TITLE = 12
_SIZE_CONTACT = 10
_SIZE_SECTION = 12
_SIZE_JOB_TITLE = 11
_SIZE_BODY = 10.5


def _set_page_size(doc: Document, region_code: str) -> None:
    """Set page dimensions and margins based on region."""
    use_letter = region_code.upper() in _LETTER_REGIONS
    section = doc.sections[0]

    if use_letter:
        section.page_width = _PAGE_WIDTH_LETTER
        section.page_height = _PAGE_HEIGHT_LETTER
    else:
        section.page_width = _PAGE_WIDTH_A4
        section.page_height = _PAGE_HEIGHT_A4

    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN


def _run_bold(para, text: str, size: float | None = None, color: RGBColor | None = None) -> None:
    """Add a bold run to a paragraph."""
    run = para.add_run(text)
    run.bold = True
    run.font.name = _FONT_BODY
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _run_normal(para, text: str, size: float | None = None) -> None:
    """Add a normal-weight run to a paragraph."""
    run = para.add_run(text)
    run.bold = False
    run.font.name = _FONT_BODY
    if size:
        run.font.size = Pt(size)


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a section heading using Heading 1 style, with a bottom border."""
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 1"]
    run = para.add_run(text.upper())
    run.font.name = _FONT_HEADING
    run.font.size = Pt(_SIZE_SECTION)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)  # near-black
    run.bold = True
    # Remove extra space above/below
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    # Add a thin bottom border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_job_entry(doc: Document, exp: dict) -> None:
    """Add a single experience entry."""
    title = exp.get("title", "")
    company = exp.get("company", "")
    location = exp.get("location", "")
    date = exp.get("date", "")
    bullets = exp.get("bullets", [])
    tech = exp.get("tech", "")

    # Job title (Heading 2 style)
    title_para = doc.add_paragraph()
    title_para.style = doc.styles["Heading 2"]
    run = title_para.add_run(title)
    run.font.name = _FONT_HEADING
    run.font.size = Pt(_SIZE_JOB_TITLE)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(0)

    # Company | Location | Date line
    meta_parts = []
    if company:
        meta_parts.append(company)
    if location:
        meta_parts.append(location)
    if date:
        meta_parts.append(date)

    if meta_parts:
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_before = Pt(0)
        meta_para.paragraph_format.space_after = Pt(2)
        run = meta_para.add_run(" | ".join(meta_parts))
        run.font.name = _FONT_BODY
        run.font.size = Pt(_SIZE_BODY)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Bullet points
    for bullet in bullets:
        if not bullet:
            continue
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.left_indent = Inches(0.25)
        run = bp.add_run(bullet)
        run.font.name = _FONT_BODY
        run.font.size = Pt(_SIZE_BODY)

    # Tech stack note
    if tech:
        tech_para = doc.add_paragraph()
        tech_para.paragraph_format.space_before = Pt(1)
        tech_para.paragraph_format.space_after = Pt(2)
        tech_run = tech_para.add_run("Stack: ")
        tech_run.font.name = _FONT_BODY
        tech_run.font.size = Pt(_SIZE_BODY)
        tech_run.bold = True
        val_run = tech_para.add_run(tech)
        val_run.font.name = _FONT_BODY
        val_run.font.size = Pt(_SIZE_BODY)


def generate_docx(cv_data: dict, region_code: str = "AU") -> bytes:
    """Generate an ATS-friendly DOCX from structured CV data.

    Returns the .docx file as bytes.
    """
    if not cv_data:
        cv_data = {}

    doc = Document()

    # Remove default styles that clash with our formatting
    _set_page_size(doc, region_code)

    # Ensure built-in styles exist (Document() ships them, but make them
    # invisible / minimal so we control all formatting explicitly)
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = _FONT_BODY

    # -------------------------------------------------------------------------
    # Header: Name + title + contact
    # -------------------------------------------------------------------------
    name = (cv_data.get("name") or "").strip()
    title = (cv_data.get("title") or "").strip()
    email = (cv_data.get("email") or "").strip()
    phone = (cv_data.get("phone") or "").strip()
    location = (cv_data.get("location") or "").strip()
    linkedin = (cv_data.get("linkedin") or "").strip()
    github = (cv_data.get("github") or "").strip()
    portfolio = (cv_data.get("portfolio") or "").strip()

    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(2)
        _run_bold(name_para, name, size=_SIZE_NAME, color=RGBColor(0x1A, 0x1A, 0x2E))

    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(3)
        run = title_para.add_run(title)
        run.font.name = _FONT_BODY
        run.font.size = Pt(_SIZE_TITLE)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Contact line: email | phone | location | linkedin
    contact_parts = []
    if email:
        contact_parts.append(email)
    if phone:
        contact_parts.append(phone)
    if location:
        contact_parts.append(location)
    if linkedin:
        contact_parts.append(linkedin)
    if github:
        contact_parts.append(github)
    if portfolio:
        contact_parts.append(portfolio)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(6)
        run = contact_para.add_run("  |  ".join(contact_parts))
        run.font.name = _FONT_BODY
        run.font.size = Pt(_SIZE_CONTACT)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # -------------------------------------------------------------------------
    # Summary
    # -------------------------------------------------------------------------
    summary = (cv_data.get("summary") or "").strip()
    if summary:
        _add_section_heading(doc, "Professional Summary")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(3)
        sp.paragraph_format.space_after = Pt(6)
        run = sp.add_run(summary)
        run.font.name = _FONT_BODY
        run.font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Experience
    # -------------------------------------------------------------------------
    experience = cv_data.get("experience") or []
    if experience:
        _add_section_heading(doc, "Experience")
        for exp in experience:
            if isinstance(exp, dict):
                _add_job_entry(doc, exp)

    # -------------------------------------------------------------------------
    # Skills
    # -------------------------------------------------------------------------
    skills = cv_data.get("skills") or []
    skills_grouped = cv_data.get("skills_grouped") or []

    if skills or skills_grouped:
        _add_section_heading(doc, "Skills")

        if skills_grouped:
            # Grouped skills: [{category, skills: [str]} or {category, skills: str}]
            for group in skills_grouped:
                if not isinstance(group, dict):
                    continue
                cat = group.get("category", "")
                grp_skills = group.get("skills", [])
                if isinstance(grp_skills, list):
                    skill_str = ", ".join(s for s in grp_skills if s)
                else:
                    skill_str = str(grp_skills)
                if not skill_str:
                    continue
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(2)
                sp.paragraph_format.space_after = Pt(1)
                if cat:
                    br = sp.add_run(f"{cat}: ")
                    br.bold = True
                    br.font.name = _FONT_BODY
                    br.font.size = Pt(_SIZE_BODY)
                vr = sp.add_run(skill_str)
                vr.font.name = _FONT_BODY
                vr.font.size = Pt(_SIZE_BODY)
        elif skills:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(3)
            sp.paragraph_format.space_after = Pt(6)
            run = sp.add_run(", ".join(str(s) for s in skills if s))
            run.font.name = _FONT_BODY
            run.font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Education
    # -------------------------------------------------------------------------
    education = cv_data.get("education") or []
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            degree = (edu.get("degree") or "").strip()
            institution = (edu.get("institution") or "").strip()
            date = (edu.get("date") or "").strip()
            if not degree:
                continue
            ep = doc.add_paragraph()
            ep.paragraph_format.space_before = Pt(4)
            ep.paragraph_format.space_after = Pt(1)
            _run_bold(ep, degree, size=_SIZE_BODY)
            meta = []
            if institution:
                meta.append(institution)
            if date:
                meta.append(date)
            if meta:
                ep.add_run("  —  " + " | ".join(meta)).font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Certifications
    # -------------------------------------------------------------------------
    certifications = cv_data.get("certifications") or []
    if certifications:
        _add_section_heading(doc, "Certifications")
        for cert in certifications:
            cert_text = cert.strip() if isinstance(cert, str) else str(cert)
            if not cert_text:
                continue
            cp = doc.add_paragraph(style="List Bullet")
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(1)
            cp.paragraph_format.left_indent = Inches(0.25)
            run = cp.add_run(cert_text)
            run.font.name = _FONT_BODY
            run.font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Projects
    # -------------------------------------------------------------------------
    projects = cv_data.get("projects") or []
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            proj_name = (proj.get("name") or proj.get("title") or "").strip()
            proj_desc = (proj.get("description") or proj.get("summary") or "").strip()
            proj_tech = (proj.get("tech") or proj.get("stack") or "").strip()
            proj_url = (proj.get("url") or proj.get("link") or "").strip()
            if not proj_name:
                continue
            pp = doc.add_paragraph()
            pp.paragraph_format.space_before = Pt(4)
            pp.paragraph_format.space_after = Pt(1)
            _run_bold(pp, proj_name, size=_SIZE_BODY)
            if proj_url:
                run = pp.add_run(f"  ({proj_url})")
                run.font.size = Pt(_SIZE_BODY)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if proj_desc:
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after = Pt(1)
                run = dp.add_run(proj_desc)
                run.font.name = _FONT_BODY
                run.font.size = Pt(_SIZE_BODY)
            if proj_tech:
                tp = doc.add_paragraph()
                tp.paragraph_format.space_before = Pt(0)
                tp.paragraph_format.space_after = Pt(2)
                br = tp.add_run("Technologies: ")
                br.bold = True
                br.font.name = _FONT_BODY
                br.font.size = Pt(_SIZE_BODY)
                vr = tp.add_run(proj_tech)
                vr.font.name = _FONT_BODY
                vr.font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Languages
    # -------------------------------------------------------------------------
    languages = cv_data.get("languages") or []
    if languages:
        _add_section_heading(doc, "Languages")
        lang_strs = []
        for lang in languages:
            if isinstance(lang, dict):
                lname = lang.get("language") or lang.get("name") or ""
                level = lang.get("level") or lang.get("proficiency") or ""
                lang_strs.append(f"{lname} ({level})" if level else lname)
            elif isinstance(lang, str):
                lang_strs.append(lang)
        if lang_strs:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(3)
            lp.paragraph_format.space_after = Pt(6)
            run = lp.add_run(", ".join(l for l in lang_strs if l))
            run.font.name = _FONT_BODY
            run.font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # References
    # -------------------------------------------------------------------------
    references = cv_data.get("references") or []
    if references:
        _add_section_heading(doc, "References")
        for ref in references:
            if not isinstance(ref, dict):
                continue
            ref_name = (ref.get("name") or "").strip()
            ref_title = (ref.get("title") or "").strip()
            ref_company = (ref.get("company") or "").strip()
            ref_contact = (ref.get("contact") or "").strip()
            if not ref_name:
                continue
            rp = doc.add_paragraph()
            rp.paragraph_format.space_before = Pt(4)
            rp.paragraph_format.space_after = Pt(1)
            _run_bold(rp, ref_name, size=_SIZE_BODY)
            meta = []
            if ref_title:
                meta.append(ref_title)
            if ref_company:
                meta.append(ref_company)
            if ref_contact:
                meta.append(ref_contact)
            if meta:
                rp.add_run("  —  " + " | ".join(meta)).font.size = Pt(_SIZE_BODY)

    # -------------------------------------------------------------------------
    # Serialise to bytes
    # -------------------------------------------------------------------------
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
